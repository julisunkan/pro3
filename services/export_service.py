import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# A4 page: 210mm wide, margins 20mm each side → 170mm usable
PAGE_W = 170


def export_email_txt(subject: str, body: str):
    """Return (content_str, mimetype, filename) for TXT export."""
    content = f"Subject: {subject}\n\n{body}"
    return content, 'text/plain', 'email.txt'


def export_email_html(subject: str, body: str):
    """Return (content_str, mimetype, filename) for HTML export."""
    body_html = (
        body
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('\n', '<br>')
    )
    subject_html = (
        subject
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
    )
    content = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>{subject_html}</title>
  <style>
    body {{ font-family: Arial, sans-serif; max-width: 600px; margin: 40px auto; color: #333; line-height: 1.6; }}
    h2 {{ border-bottom: 1px solid #eee; padding-bottom: 8px; }}
  </style>
</head>
<body>
  <h2>{subject_html}</h2>
  <p>{body_html}</p>
</body>
</html>"""
    return content, 'text/html', 'email.html'


def _safe(text: str) -> str:
    """Encode to latin-1, replacing unsupported characters, and strip long runs."""
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _wrap_line(line: str, max_chars: int = 95) -> list:
    """
    Break a single line into chunks of at most max_chars characters.
    Preserves natural word boundaries where possible.
    """
    if len(line) <= max_chars:
        return [line]

    words = line.split(' ')
    chunks = []
    current = ''
    for word in words:
        # If the word itself is longer than max_chars, hard-break it
        while len(word) > max_chars:
            chunks.append(word[:max_chars])
            word = word[max_chars:]

        candidate = f'{current} {word}'.strip() if current else word
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = word

    if current:
        chunks.append(current)
    return chunks


def export_email_pdf(subject: str, body: str):
    """Return (bytes, mimetype, filename) for PDF export."""
    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # ── Subject ──
    pdf.set_font('Helvetica', 'B', 13)
    pdf.set_x(20)
    pdf.cell(PAGE_W, 8, _safe(f'Subject: {subject}'), ln=True)
    pdf.ln(2)

    # ── Divider ──
    pdf.set_draw_color(180, 180, 180)
    y = pdf.get_y()
    pdf.line(20, y, 190, y)
    pdf.ln(5)

    # ── Body ──
    pdf.set_font('Helvetica', '', 11)
    for raw_line in body.split('\n'):
        safe_line = _safe(raw_line)
        if safe_line.strip():
            for chunk in _wrap_line(safe_line):
                pdf.set_x(20)
                pdf.cell(PAGE_W, 6, chunk, ln=True)
        else:
            pdf.ln(4)   # blank paragraph spacing

    pdf_bytes = pdf.output()
    return bytes(pdf_bytes), 'application/pdf', 'email.pdf'


def export_email_docx(subject: str, body: str):
    """Return (bytes, mimetype, filename) for DOCX export."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(72)
        section.right_margin  = Pt(72)

    # Subject heading
    subj_para = doc.add_paragraph()
    subj_run  = subj_para.add_run(f"Subject: {subject}")
    subj_run.bold      = True
    subj_run.font.size = Pt(14)
    subj_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Divider
    doc.add_paragraph('─' * 60)

    # Body paragraphs
    for line in body.split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return buf.read(), mime, 'email.docx'
